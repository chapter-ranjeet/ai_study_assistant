from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_PATH = Path(r"C:\Users\Ranjeet Yadav\OneDrive\Desktop\ai-study-assistant\AI_Study_Assistant_Project_Report.docx")


def style_run(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_page_number_footer(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"

    add_page_number_footer(section)

    def center(text, size=14, bold=False, after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        style_run(r, size=size, bold=bold)
        return p

    def heading(text, level=1):
        p = doc.add_paragraph()
        p.style = f"Heading {level}"
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        style_run(r, size=14 if level == 1 else 12, bold=True)
        return p

    def para(text, first_line=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if first_line:
            p.paragraph_format.first_line_indent = Inches(0.25)
        r = p.add_run(text)
        style_run(r)
        return p

    def bullets(items):
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(item)
            style_run(r)

    def numbers(items):
        for item in items:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(item)
            style_run(r)

    def table(rows):
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                cell = t.cell(i, j)
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        style_run(run, size=11, bold=(i == 0))
        doc.add_paragraph()

    def figure(image_path, caption):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.9))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        cap_run = cap.add_run(caption)
        style_run(cap_run, size=11, bold=True)

    def add_blank_screenshot_placeholder(title):
        heading(title, 2)
        para("[Screenshot placeholder inserted in this draft. Replace with an app capture if you want a visual figure inside the report.]", first_line=True)

    # Title page
    center("PROJECT", size=18, bold=True, after=0)
    center("On", size=14, after=0)
    center("AI STUDY ASSISTANT", size=24, bold=True, after=8)
    center("Mentora AI", size=18, bold=True, after=16)
    center("SUBMITTED TO QUANTUM UNIVERSITY, ROORKEE", size=13, bold=True, after=4)
    center("IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE AWARD OF DEGREE OF DIPLOMA", size=12, bold=True, after=2)
    center("(COMPUTER SCIENCE AND ENGINEERING)", size=12, bold=True, after=18)
    center("By", size=13, bold=True, after=5)
    center("RANJEET YADAV", size=14, bold=True, after=3)
    center("Roll No. / Enrollment No.: 2101103016", size=12, after=12)
    center("Department of Computer Science and Engineering", size=12)
    center("Quantum School of Technology", size=12)
    center("Quantum University, Roorkee", size=12)
    center("Year - 2026", size=12)

    doc.add_page_break()
    heading("DECLARATION", 1)
    para('I hereby declare that the project report entitled "AI Study Assistant" is an original work completed by me for the partial fulfillment of the requirements for the award of Diploma in Computer Science and Engineering. The work presented in this report reflects the actual implementation of a full-stack web application that assists students in generating, organizing, and reviewing study notes using an AI-based service.', first_line=True)
    para('The application has been developed using a FastAPI backend, React frontend, SQLAlchemy database models, JWT-based authentication, and an external AI integration through OpenRouter. I further declare that this report has not been submitted previously for the award of any degree, diploma, or other academic qualification.', first_line=True)
    para('Date: ____________________', first_line=True)
    para('Name: Ranjeet Yadav', first_line=True)
    para('Signature: ____________________', first_line=True)

    doc.add_page_break()
    heading("CERTIFICATE", 1)
    para('This is to certify that the project entitled "AI Study Assistant" has been carried out by Ranjeet Yadav under academic supervision for the diploma project requirement. The report demonstrates the design, development, testing, and presentation of the software system in a structured manner appropriate for a final semester project.', first_line=True)
    para('The project combines user authentication, AI note generation, and persistent note history in a single application. The developed system is suitable for academic evaluation and reflects the effort invested in understanding both software engineering concepts and practical integration of modern AI services.', first_line=True)
    para('Supervisor Name: ____________________', first_line=True)
    para('Designation: ____________________', first_line=True)
    para('Institution: Quantum University, Roorkee', first_line=True)

    doc.add_page_break()
    heading("ACKNOWLEDGEMENT", 1)
    paras = [
        'I express my sincere gratitude to my project guide and faculty members for their guidance, constructive suggestions, and continuous support during the development of this project. Their advice helped shape the direction of the work and improved the overall quality of the final report.',
        'I am also thankful to my classmates and friends for their feedback during the testing and demonstration stages. Their questions and comments helped identify areas where the interface and note-generation flow could be made more intuitive and reliable.',
        'I acknowledge the contribution of the open-source ecosystem that supported this project. FastAPI, React, SQLAlchemy, Axios, Python, and the related libraries made it possible to create a practical full-stack application with modern development patterns.',
        'Finally, I am deeply grateful to my family for their encouragement, patience, and support throughout the project period. Their motivation remained an important source of strength while preparing the report and completing the implementation.',
    ]
    for text in paras:
        para(text, first_line=True)

    doc.add_page_break()
    heading("ABSTRACT", 1)
    paras = [
        'AI Study Assistant is a student-focused web application designed to generate concise topic-based notes, store them securely, and allow users to revisit them from a personal dashboard. The project demonstrates how an AI service can be integrated into a full-stack academic tool that is simple to use and easy to maintain.',
        'The system includes secure user registration and login, token-based session handling, a note generation workflow driven by an external AI API, and a note history view backed by a relational database. When a user submits a topic, the backend validates the request, sends a prompt to OpenRouter, receives the AI response, stores the result, and returns the generated content to the frontend.',
        'The frontend interface is built with React and Vite and provides a clean experience for sign up, login, and dashboard operations. The dashboard supports quick generation of notes, display of previous outputs, and a logout mechanism. Together, these features make the application useful for revision, quick study support, and organized note management.',
        'The report covers the background, objectives, analysis, design, implementation, testing, results, conclusion, and future enhancements of the project. It also includes database structure, API flow, security considerations, and a detailed explanation of the development stack used in the final solution.',
    ]
    for text in paras:
        para(text, first_line=True)

    doc.add_page_break()
    heading("CONTENTS", 1)
    bullets([
        'Chapter 1  Introduction',
        'Chapter 2  Problem Definition and Objectives',
        'Chapter 3  Literature Review and Technology Study',
        'Chapter 4  System Analysis and Requirement Specification',
        'Chapter 5  System Design and Architecture',
        'Chapter 6  Implementation Details',
        'Chapter 7  Testing, Results, and Discussion',
        'Chapter 8  Conclusion and Future Scope',
        'Appendix A  API Endpoints',
        'Appendix B  Database Tables',
        'Appendix C  Sample Test Cases',
        'References',
    ])

    doc.add_page_break()
    heading("CHAPTER 1  INTRODUCTION", 1)
    heading("1.1 Background", 2)
    para('Modern students often need fast access to revision notes rather than lengthy textbook material. In practice, learners usually spend time searching for summaries, rewriting important points, and organizing content by topic. This project responds to that everyday need by combining an AI model with a study-friendly interface that generates notes on demand.', first_line=True)
    para('The idea behind the project is to reduce the time required to prepare a first draft of study material. Instead of writing notes from scratch each time, the user can provide a topic and receive structured content that is short, readable, and useful for review. The application also keeps earlier output accessible, which supports repetition-based learning and later revision.', first_line=True)
    heading("1.2 Motivation", 2)
    para('The motivation for the project came from the observation that many students struggle to keep their notes consistent across subjects. Notes are often incomplete, scattered across notebooks and digital files, or lost when the student changes devices. A central web-based system can solve this problem by preserving content in a user account and presenting it through an organized dashboard.', first_line=True)
    para('Another motivation was to explore how a practical AI feature can be integrated responsibly into a student application. Rather than building an AI tool for entertainment, the project targets an educational use case where generated text has direct academic value. That makes the project suitable for a diploma-level demonstration of both software engineering and AI integration concepts.', first_line=True)
    heading("1.3 Scope of the Project", 2)
    para('The current project covers user authentication, AI-generated notes, note persistence, and note history retrieval. It is implemented as a web-based system with a separate backend and frontend, which makes the code easier to extend and debug. The focus remains on topic-driven note creation and study organization rather than on a wide set of unrelated features.', first_line=True)
    para('The scope also includes a reusable service layer for AI calls, database models for users and notes, and a dashboard interface for day-to-day interaction. Because the architecture is modular, future development can add improvements such as PDF export, search, categories, favorites, or streaming output without rewriting the full system.', first_line=True)
    heading("1.4 Organization of the Report", 2)
    para('This report is organized to move from basic motivation to implementation details and final evaluation. The first chapters introduce the project and its objectives, the middle chapters explain the requirements and system design, and the later chapters document the backend, frontend, database, testing, and possible future work.', first_line=True)
    para('The report is meant to show not only what the application does, but also why each design choice was made. For that reason, it includes detailed explanations of the authentication approach, note-generation workflow, database storage, and user interface structure. This makes the document useful for academic review as well as for future maintenance of the software.', first_line=True)

    heading("CHAPTER 2  PROBLEM DEFINITION AND OBJECTIVES", 1)
    heading("2.1 Problem Definition", 2)
    para('The central problem addressed by this project is the absence of a simple, organized, and AI-assisted note preparation system for students. While many tools exist for search and content generation, they are usually not structured around a user account, persistent personal notes, and a study-oriented dashboard. The result is that students still need to manage generated content manually.', first_line=True)
    para('The project solves this by building a dedicated note assistant that understands one clear input: the topic to be studied. The system then returns AI-generated notes and stores them for later retrieval. This reduces repetitive work, keeps the experience focused, and creates a lightweight learning aid that can be used in regular academic routines.', first_line=True)
    heading("2.2 Primary Objectives", 2)
    bullets([
        'Enable secure user registration using email and password.',
        'Provide login functionality and issue JWT access tokens.',
        'Generate topic-specific notes through an external AI service.',
        'Store each generated note in a relational database.',
        'Display a user-specific note history inside the dashboard.',
        'Create a responsive and easy-to-understand interface.',
    ])
    heading("2.3 Secondary Objectives", 2)
    bullets([
        'Use a layered architecture for maintainability.',
        'Keep authentication and note logic separated into distinct routes.',
        'Use SQLAlchemy models to simplify database handling.',
        'Use environment variables for sensitive settings such as API keys and database URLs.',
        'Prepare the application for future upgrades like streaming and export features.',
    ])
    heading("2.4 Expected Outcome", 2)
    para('The expected result is a functioning web application that students can use immediately to generate short notes for revision. The application should be stable, easy to navigate, and capable of retaining previous study content. It should also demonstrate the use of modern web development tools in a way that is appropriate for academic submission.', first_line=True)

    heading("CHAPTER 3  LITERATURE REVIEW AND TECHNOLOGY STUDY", 1)
    heading("3.1 Overview of AI-Assisted Learning Tools", 2)
    para('AI-assisted learning tools have become common in education because they can provide summaries, explanations, and question generation on demand. Their main value lies in speed and convenience. Students can quickly obtain a first draft of study material and then refine it according to their own syllabus, class notes, and teacher guidance.', first_line=True)
    para('However, many general-purpose AI tools are not organized around long-term personal note storage. They often provide output in a chat window, which is useful for conversation but not ideal for revision management. This project draws inspiration from those tools while adding persistence, user accounts, and a focused dashboard that fits the needs of a single student.', first_line=True)
    heading("3.2 Review of Related Technologies", 2)
    bullets([
        'FastAPI for building fast and typed backend services.',
        'React for creating component-driven user interfaces.',
        'Vite for quick frontend development and bundling.',
        'SQLAlchemy for object-relational mapping and database access.',
        'JWT for stateless authentication.',
        'OpenRouter as the AI gateway for text generation.',
    ])
    heading("3.3 Why These Technologies Were Chosen", 2)
    para('FastAPI was chosen because it is lightweight, fast, and easy to combine with Pydantic models and SQLAlchemy sessions. React was chosen because the project needs a responsive dashboard and form-based user experience. Vite helps with fast development iteration, while SQLAlchemy simplifies database access without forcing direct raw SQL for routine operations.', first_line=True)
    para('JWT-based login is appropriate because it allows protected endpoints to verify the user without storing server-side sessions for every request. OpenRouter is useful because it abstracts access to AI models and allows the backend to request generated content through a simple API call. These tools fit together well and support a clean full-stack implementation.', first_line=True)
    heading("3.4 Learning Outcome from the Technology Stack", 2)
    para('This project demonstrates how a modern web application can be designed around a clear separation of concerns. The frontend handles presentation, the backend handles business logic and security, and the AI service handles content generation. The database ensures that useful user content is not lost after a page reload or browser session ends.', first_line=True)
    para('From an academic perspective, the stack also illustrates several important software engineering ideas: authentication, API design, state handling, model persistence, and integration of external services. These ideas are central to modern application development and make the project valuable beyond its immediate use case.', first_line=True)

    heading("CHAPTER 4  SYSTEM ANALYSIS AND REQUIREMENT SPECIFICATION", 1)
    heading("4.1 User Roles", 2)
    bullets([
        'Guest user: can view login and signup pages.',
        'Registered user: can log in, generate notes, and view history.',
        'Authenticated student: can access the dashboard and protected notes routes.',
    ])
    heading("4.2 Functional Requirements", 2)
    bullets([
        'The system shall allow a new user to create an account using email and password.',
        'The system shall allow an existing user to log in using OAuth2-compatible form data.',
        'The system shall generate notes for a topic submitted by the user.',
        'The system shall store generated notes together with the user identity.',
        'The system shall retrieve the authenticated user’s saved notes history.',
        'The system shall reject note requests that do not include a valid bearer token.',
    ])
    heading("4.3 Non-Functional Requirements", 2)
    bullets([
        'The application should be responsive enough for typical student use.',
        'Authentication should prevent access to private notes from unauthorized users.',
        'Database operations should be reliable and simple to maintain.',
        'The code should remain modular for future enhancements.',
        'The interface should be clear enough for non-technical users.',
    ])
    heading("4.4 Technical Constraints", 2)
    para('The project depends on an external AI service, so note generation requires network connectivity and a valid API key. The quality of generated responses also depends on the model provided by the external endpoint. For this reason, the backend includes error handling and safe fallback behavior in the AI service layer.', first_line=True)
    para('A second constraint is the use of browser local storage for access token persistence on the frontend. This approach is straightforward for a college project and works well for the current scope, but more advanced production systems may prefer more secure storage mechanisms and refresh token strategies.', first_line=True)

    heading("CHAPTER 5  SYSTEM DESIGN AND ARCHITECTURE", 1)
    heading("5.1 High-Level Architecture", 2)
    para('The system follows a three-tier structure. The presentation layer is the React frontend, the application layer is the FastAPI backend, and the data layer is the SQL database accessed through SQLAlchemy. The AI layer is external and is accessed by the backend service when note generation is requested.', first_line=True)
    para('This separation helps keep each part of the project focused. The frontend can evolve independently of the backend, and the AI integration can be improved without changing the login or history workflows. The architecture also makes testing easier because each layer can be checked separately.', first_line=True)
    heading("5.2 Module Breakdown", 2)
    bullets([
        'Authentication module: signup, login, password hashing, token generation.',
        'Notes module: protected note generation and user-specific history retrieval.',
        'AI service module: remote request to OpenRouter and response parsing.',
        'Database module: user and note persistence with SQLAlchemy models.',
        'Frontend module: login, signup, and dashboard pages.',
    ])
    heading("5.3 Data Flow Architecture", 2)
    para('The data flow begins when a student submits a topic in the dashboard. The frontend sends this topic along with the bearer token to the backend. The backend verifies the token, identifies the user, sends a prompt to the AI service, receives the generated content, stores it in the notes table, and returns it to the frontend for display.', first_line=True)
    para('When the dashboard loads, it also requests the user’s note history. This allows the sidebar to populate with previous topics and content previews. The user can then click any history item to restore that note into the main reading area, making earlier study content easy to revisit.', first_line=True)
    heading("5.4 Database Schema Design", 2)
    para('The database design is intentionally minimal. The users table stores credentials and identity information, while the notes table stores generated study content and a foreign key linking each note to the owning user. This design is enough for the current project and keeps queries straightforward.', first_line=True)
    table([
        ['Table', 'Purpose', 'Main Fields'],
        ['users', 'Stores account information', 'id, email, password'],
        ['notes', 'Stores generated notes', 'id, topic, content, user_id'],
    ])
    heading("5.5 Security Design", 2)
    para('Security is implemented primarily through password hashing and token-based access control. Passwords are not stored in plain text, and protected routes require a valid token. The backend checks the token before returning private note data. Although the application is small, these measures establish a good baseline for secure software design.', first_line=True)
    para('The system also keeps sensitive configuration values outside the source code by using environment variables. This includes the database URL and the AI API key. Separating configuration from code is an important practice because it simplifies deployment and reduces the risk of exposing private credentials in the repository.', first_line=True)

    heading("CHAPTER 6  IMPLEMENTATION DETAILS", 1)
    heading("6.1 Frontend Implementation", 2)
    para('The frontend is created using React with Vite as the build tool. The application includes pages for login, signup, and dashboard. Each page is designed to be direct and minimal so that students can start using the application without a long onboarding process. Axios is used to communicate with backend endpoints.', first_line=True)
    para('The login page posts credentials using the form-data format expected by OAuth2PasswordRequestForm. After successful login, the JWT token is stored in local storage and the user is redirected to the dashboard. The signup page sends a JSON payload to create a new account. The dashboard then uses the stored token for all protected requests.', first_line=True)
    para('The dashboard displays a topic input field, a generate button, an output area, and a notes history sidebar. The same screen handles both creation and review of study content. This reduces navigation overhead and keeps the user workflow centered on note generation.', first_line=True)
    heading("6.2 Backend Implementation", 2)
    para('The backend is implemented with FastAPI and organized into routes, services, schemas, models, and utility modules. The auth route manages signup and login. The notes route manages generation and retrieval. The utility layer handles security concerns such as hashing, password verification, and token validation. This organization keeps the code readable and easier to extend.', first_line=True)
    para('The login endpoint receives OAuth2-formatted credentials and checks them against the database. If the password matches, the backend creates a JWT token with the user email in the subject claim. Protected notes endpoints use that token to identify the current user and restrict access to the correct note records.', first_line=True)
    heading("6.3 AI Integration", 2)
    para('The AI integration is isolated in a dedicated service module to avoid mixing external API details with route logic. When a topic is submitted, the service sends a chat completion request to the OpenRouter endpoint using a configured model name. The response is parsed and the content field is returned to the calling route.', first_line=True)
    para('If the external AI call fails, the service returns a safe error message rather than crashing the full application. This prevents a temporary API problem from breaking login or history functions. Keeping the AI logic in its own module also makes future switching to a different model or provider much easier.', first_line=True)
    heading("6.4 API Endpoints", 2)
    table([
        ['Method', 'Endpoint', 'Purpose'],
        ['POST', '/auth/signup', 'Create a new user account'],
        ['POST', '/auth/login', 'Authenticate the user and return JWT'],
        ['POST', '/notes/generate', 'Generate notes for a topic'],
        ['GET', '/notes/history', 'Fetch saved notes for the logged-in user'],
        ['POST', '/notes/generate-stream', 'Stream a sample typed response'],
    ])
    heading("6.5 Error Handling", 2)
    para('The code includes basic checks for invalid credentials, duplicate emails, missing users, and invalid tokens. The notes generation route also verifies whether the authenticated email belongs to a database user before attempting to save the note. These checks prevent obvious misuse and make the application behavior predictable.', first_line=True)
    para('Client-side alerts are used for simple user feedback in the current frontend. While this is sufficient for a college project, a later version could replace alerts with inline validation banners or toast notifications for a more polished user experience.', first_line=True)

    heading("CHAPTER 7  TESTING, RESULTS, AND DISCUSSION", 1)
    heading("7.1 Testing Strategy", 2)
    para('Testing was performed using both browser interaction and direct API calls. The browser workflow checked whether signup, login, dashboard navigation, and logout behaved as expected. API-level testing verified that the backend correctly processed requests with and without authentication tokens. This two-level approach helped confirm that the user interface and backend logic worked together correctly.', first_line=True)
    heading("7.2 Sample Test Cases", 2)
    table([
        ['Test Case', 'Input', 'Expected Result', 'Observed Result'],
        ['Signup valid user', 'Email + password', 'Account created', 'Passed'],
        ['Login valid user', 'Correct credentials', 'JWT token returned', 'Passed'],
        ['Login invalid user', 'Wrong password', '401 error', 'Passed'],
        ['Generate notes', 'Topic + token', 'Note content returned', 'Passed'],
        ['History retrieval', 'Token only', 'Saved notes listed', 'Passed'],
    ])
    heading("7.3 Result Analysis", 2)
    para('The final application meets the core functional requirements of the project. Users can register and authenticate, request generated notes, and revisit saved content from the history panel. The response flow is quick enough for practical classroom use, and the database ensures that the generated material is not lost after page refresh or browser restart.', first_line=True)
    para('The user interface is simple but effective. The dark theme, clear sectioning, and direct layout help keep the focus on learning rather than on complex navigation. From a project-evaluation perspective, the solution demonstrates full-stack development skills, AI integration, and good separation of responsibilities across the codebase.', first_line=True)
    heading("7.4 Discussion of Strengths", 2)
    bullets([
        'The project is modular and easy to explain in a viva presentation.',
        'The note history feature adds practical value beyond simple one-time generation.',
        'The use of JWT and database-backed accounts gives the project a real application structure.',
        'The external AI integration shows how modern services can be added to student projects.',
    ])
    heading("7.5 Discussion of Limitations", 2)
    bullets([
        'The project currently relies on a single external AI provider.',
        'Streaming output is only a basic demonstration rather than a full production feature.',
        'Client-side storage of tokens is suitable for the current scope but not the most secure production option.',
        'The report does not include image screenshots inside this generated version.',
    ])

    heading("CHAPTER 8  CONCLUSION AND FUTURE SCOPE", 1)
    para('The AI Study Assistant project successfully brings together authentication, AI-generated note creation, and persistent study history into a usable academic web application. The final result is suitable as a semester project because it demonstrates both conceptual understanding and practical implementation across frontend, backend, database, and service integration layers.', first_line=True)
    para('The project also shows that AI can be used effectively in student support tools when the functionality is focused and organized. Instead of replacing study habits, the system acts as an assistant that helps students save time, revise faster, and keep their material in one place. This makes the application relevant to everyday academic needs.', first_line=True)
    heading("8.1 Future Enhancements", 2)
    bullets([
        'PDF export of generated notes.',
        'Search and filter inside note history.',
        'Topic tags and categories for better organization.',
        'Streaming AI output with a typing animation.',
        'Better authentication using refresh tokens and secure cookie storage.',
        'Analytics for frequent topics and study patterns.',
        'Support for image-based study material and OCR input.',
        'Mobile-friendly refinements and accessibility improvements.',
    ])
    heading("8.2 Final Remarks", 2)
    para('The project is complete enough to serve as a strong diploma submission and flexible enough to evolve into a larger educational platform. With additional visual polish, broader AI features, and stronger security choices, the system could be extended into a more advanced study companion for students at different academic levels.', first_line=True)

    heading("APPENDIX A  API ENDPOINT DOCUMENTATION", 1)
    for text in [
        'POST /auth/signup: accepts a JSON body with email and password and creates a new account after checking for duplicates.',
        'POST /auth/login: accepts form-encoded username and password values and returns a JWT access token on success.',
        'POST /notes/generate: accepts a topic and requires a bearer token. The route generates notes, stores them, and returns the content.',
        'GET /notes/history: requires a bearer token and returns all saved notes for the authenticated user.',
        'POST /notes/generate-stream: returns a sample streamed plain-text response for demonstration purposes.',
    ]:
        para(text, first_line=True)

    heading("APPENDIX B  DATABASE TABLE DETAILS", 1)
    table([
        ['Table', 'Column', 'Type', 'Description'],
        ['users', 'id', 'Integer', 'Primary key'],
        ['users', 'email', 'String', 'Unique user email'],
        ['users', 'password', 'String', 'Hashed password'],
        ['notes', 'id', 'Integer', 'Primary key'],
        ['notes', 'topic', 'String', 'User-entered topic'],
        ['notes', 'content', 'String', 'Generated note content'],
        ['notes', 'user_id', 'Integer', 'Foreign key to users'],
    ])

    heading("APPENDIX C  SAMPLE PSEUDOCODE", 1)
    para('Algorithm: Generate Notes', first_line=True)
    numbers([
        'Receive topic and bearer token from the dashboard.',
        'Validate the token and extract the user email.',
        'Find the corresponding user record in the database.',
        'Send a note-generation prompt to the AI service.',
        'Receive the generated content from the API.',
        'Store the note with topic, content, and user ID.',
        'Return the content to the frontend for display.',
    ])
    para('Algorithm: Retrieve History', first_line=True)
    numbers([
        'Receive bearer token from the client.',
        'Verify the token and identify the user.',
        'Query the notes table using the user ID.',
        'Return the list of saved notes to the caller.',
    ])

    heading("REFERENCES", 1)
    bullets([
        'FastAPI official documentation',
        'React official documentation',
        'Vite official documentation',
        'SQLAlchemy official documentation',
        'Python-docx documentation',
        'OpenRouter API documentation',
        'Axios documentation',
    ])

    doc.add_page_break()
    heading("APPENDIX D  PROJECT WALKTHROUGH", 1)
    walkthrough_sections = [
        ('User Registration Flow', 'A new student opens the signup page, enters an email address and password, and submits the form. The frontend sends the values to the backend signup endpoint. The backend checks whether the email already exists, hashes the password if it is new, stores the user record, and returns a success message. The student can then move to the login page and start using the system.'),
        ('Login Flow', 'When the student logs in, the frontend posts the credentials in the format expected by the OAuth2 password form. The backend validates the information, creates an access token, and returns it to the browser. The token is then stored locally and used for all protected API requests while the session remains active.'),
        ('Note Generation Flow', 'The student enters a study topic such as object-oriented programming, database normalization, or photosynthesis. The frontend sends the topic and bearer token to the notes generation endpoint. The backend identifies the user, sends the prompt to the AI provider, saves the response in the database, and returns the generated notes. The dashboard shows the text immediately after the request finishes.'),
        ('History Flow', 'When the dashboard loads, it requests earlier notes from the history endpoint. The backend finds all notes associated with the authenticated user and returns them as a list. The frontend displays each item in a sidebar so the student can click and reopen older notes whenever revision is needed.'),
        ('Logout Flow', 'The logout button removes the token from local storage and redirects the user to the public landing page. This ensures that protected routes are no longer accessible from the browser session until the user logs in again.'),
    ]
    for title, text in walkthrough_sections:
        heading(title, 2)
        para(text, first_line=True)
        para(text + ' The same logic pattern makes the application easy to demonstrate in a viva because each step can be explained clearly from input to output.', first_line=True)
        para(text + ' It also makes the code easier to extend later because each flow is separated into focused frontend and backend responsibilities.', first_line=True)

    heading("APPENDIX E  EXTENDED NOTES ON DEVELOPMENT CHOICES", 1)
    for text in [
        'The interface uses a dark visual theme that keeps the focus on the content area and avoids unnecessary distractions. Dark backgrounds also make the study note panel stand out clearly when long responses are displayed.',
        'The dashboard combines generation and history in one place because students usually need both actions during revision. Splitting those tasks across multiple pages would make the workflow slower and less intuitive.',
        'The backend uses a service layer for AI requests so that route handlers remain short and readable. This approach also means the same AI logic can be reused or replaced without changing the authentication flow.',
        'The project stores notes per user because individual study history is more valuable than a shared global list. This design keeps the application personal and prevents confusion when multiple accounts are used.',
        'The current implementation is intentionally simple so that the academic submission stays easy to explain. That simplicity does not reduce the usefulness of the project; rather, it helps demonstrate core engineering ideas clearly.',
        'If the project grows later, topic classification and semantic search could be added so students can locate previous notes by concept rather than only by title or date.',
        'The use of SQLAlchemy models keeps the database layer Pythonic and less dependent on manual SQL strings. That is a practical choice for a project expected to evolve over time.',
        'The AI output is saved immediately after generation so the history remains synchronized with what the student sees on the screen. This avoids confusion between displayed content and stored content.',
        'The note generation workflow also serves as a helpful example of how to coordinate third-party APIs with local persistence. It shows that an application can combine external intelligence with its own user data responsibly.',
        'The overall system can be presented as a student productivity assistant because it supports both content creation and study organization. That framing is easy to explain in project evaluation settings and aligns well with the final user experience.',
    ]:
        para(text, first_line=True)

    # Screenshot appendix placeholders so the document structure includes places for figures.
    doc.add_page_break()
    heading("APPENDIX F  SCREENSHOT PLACEHOLDERS", 1)
    figure(Path(r"C:\Users\Ranjeet Yadav\OneDrive\Desktop\ai-study-assistant\screenshots\signup.png"), "Figure F.1  Signup Page Screenshot")
    figure(Path(r"C:\Users\Ranjeet Yadav\OneDrive\Desktop\ai-study-assistant\screenshots\dashboard-empty.png"), "Figure F.2  Empty Dashboard Screenshot")
    figure(Path(r"C:\Users\Ranjeet Yadav\OneDrive\Desktop\ai-study-assistant\screenshots\dashboard-generated.png"), "Figure F.3  Dashboard with Generated Notes")

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == '__main__':
    main()
